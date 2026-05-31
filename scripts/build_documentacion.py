"""
Genera el contenido de Documentacion.docx conservando la carátula existente.
Ejecutar: python scripts/build_documentacion.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Documentacion.docx"
COVER_END_MARKER = "Guatemala, 30 de mayo de 2026"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, f"• {item}")


def trim_after_cover(doc: Document) -> None:
    """Elimina contenido previo tras la carátula (si se re-ejecuta el script)."""
    found = False
    to_remove = []
    for i, para in enumerate(doc.paragraphs):
        if COVER_END_MARKER in (para.text or ""):
            found = True
            continue
        if found and para.text.strip():
            to_remove.append(para)
    for para in to_remove:
        p = para._element
        p.getparent().remove(p)


def append_documentation(doc: Document) -> None:
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Call Center Inteligente — La Casa del Sabor")
    tr.bold = True
    tr.font.size = Pt(16)
    add_para(
        doc,
        "Documentación técnica del proyecto final. Sistema web de pedidos por llamada "
        "con agente de voz basado en IA, recuperación semántica sobre menú y paneles "
        "operativos en tiempo real.",
    )

    # 1
    add_heading(doc, "1. Descripción del problema", 1)
    add_para(
        doc,
        "Los restaurantes que reciben pedidos por teléfono dependen de operadores humanos "
        "para consultar el menú, anotar platillos, confirmar datos (nombre, tipo de pedido, "
        "dirección, forma de pago) y comunicar el pedido a cocina y caja. Este proceso es "
        "lento, propenso a errores (platillos inexistentes, precios incorrectos, órdenes "
        "incompletas) y difícil de auditar en tiempo real.",
    )
    add_para(
        doc,
        "El proyecto aborda el caso de uso del Grupo 2 del curso: un call center inteligente "
        "que atiende llamadas (reales o simuladas), interpreta la intención del cliente, "
        "consulta un menú mediante búsqueda semántica (RAG), estructura el pedido con salida "
        "tipada (structured output) y refleja el estado del pedido en módulos operativos "
        "(cocina, caja, pedidos activos), de forma similar a un restaurante real.",
    )

    # 2
    add_heading(doc, "2. Objetivo del sistema", 1)
    add_para(
        doc,
        "Diseñar e implementar una aplicación web que permita:",
    )
    add_bullets(
        doc,
        [
            "Simular o recibir llamadas de clientes que desean ordenar en el restaurante «La Casa del Sabor».",
            "Transcribir y sintetizar voz (STT/TTS) cuando el canal lo requiera.",
            "Conducir la conversación con un agente LLM apoyado en herramientas y RAG sobre el menú.",
            "Convertir la conversación en una orden estructurada (ítems, totales, datos de entrega y pago).",
            "Mostrar y actualizar el estado del pedido en tiempo real en paneles de cocina, caja y listado general.",
            "Cumplir los requisitos académicos del curso: LLM, tools, RAG, embeddings, vector store, workflow agentic y persistencia.",
        ],
    )

    # 3
    add_heading(doc, "3. Funcionalidades principales", 1)

    add_heading(doc, "3.1 Canal de atención (Agente)", 2)
    add_bullets(
        doc,
        [
            "Opción 1 — Simulación web: chat por texto o micrófono (Whisper + TTS) sin depender de telefonía.",
            "Opción 2 — Teléfono (Twilio): llamada al número configurado; monitor en vivo en la web (requiere túnel público).",
            "Memoria conversacional por sesión (transcript persistido entre turnos).",
            "Visualización del estado del pedido y barra de flujo (recibido → cocina → entregado).",
            "Botón «Confirmar y enviar a cocina» cuando el cliente confirma el pedido.",
        ],
    )

    add_heading(doc, "3.2 Inteligencia artificial", 2)
    add_bullets(
        doc,
        [
            "Agente LangGraph: nodo agente → herramientas → nodo de salida estructurada.",
            "Tools: buscar_en_menu (RAG), validar_disponibilidad, calcular_total_pedido.",
            "RAG con ChromaDB y embeddings OpenAI sobre menu.json.",
            "Structured output: modelos Pydantic StructuredOrder y AgentResponse.",
        ],
    )

    add_heading(doc, "3.3 Operaciones del restaurante", 2)
    add_bullets(
        doc,
        [
            "Panel Cocina: pedidos confirmados, en preparación y listos.",
            "Panel Caja: pedidos listos para cobro/entrega y pedidos en proceso.",
            "Panel Todos los pedidos: vista completa con avance de estados.",
            "WebSocket /ws/orders para actualización en tiempo real entre pestañas.",
        ],
    )

    add_heading(doc, "3.4 Flujo de estados del pedido", 2)
    add_para(
        doc,
        "Recibido → Confirmando → Confirmado / En cocina → Listo → Entregado. "
        "Al confirmar (frases del cliente o botón), el pedido pasa a cocina automáticamente.",
    )

    # 4
    add_heading(doc, "4. Tecnologías utilizadas", 1)

    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    hdr[0].text = "Capa"
    hdr[1].text = "Tecnología"
    rows = [
        ("Frontend", "React 18, Vite, JavaScript, WebSocket nativo"),
        ("Backend", "Python 3, FastAPI, Uvicorn, SQLAlchemy (async), SQLite"),
        ("IA / Agentes", "LangChain, LangGraph, OpenAI GPT-4o-mini"),
        ("RAG", "OpenAI Embeddings (text-embedding-3-small), ChromaDB"),
        ("Voz", "OpenAI Whisper (STT), OpenAI TTS"),
        ("Telefonía (opcional)", "Twilio Programmable Voice, TwiML"),
        ("Tiempo real", "WebSockets (FastAPI)"),
        ("Validación / API", "Pydantic v2"),
    ]
    for a, b in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    doc.add_paragraph()

    # 5
    add_heading(doc, "5. Explicación del flujo general", 1)

    add_heading(doc, "5.1 Flujo de una llamada simulada", 2)
    add_bullets(
        doc,
        [
            "El agente inicia sesión (POST /api/calls/simulation/start) y guarda transcript en SQLite.",
            "Cada turno del cliente (texto o audio) se transcribe si aplica y se envía al grafo LangGraph.",
            "El agente puede invocar tools (menú RAG, validación, total) antes de generar respuesta.",
            "El nodo structure produce AgentResponse con spoken_response y StructuredOrder.",
            "Se persiste/actualiza la orden y se notifica por WebSocket a Cocina y Caja.",
            "Al confirmar, el estado pasa a en_cocina; operadores avanzan estados manualmente en los paneles.",
        ],
    )

    add_heading(doc, "5.2 Flujo de datos (IA)", 2)
    add_para(
        doc,
        "Cliente → STT (opcional) → LangGraph (LLM + tools + RAG) → Structured Output → "
        "SQLite → WebSocket → UI (Agente, Cocina, Caja).",
    )

    # 6
    add_heading(doc, "6. Decisiones técnicas importantes", 1)
    add_bullets(
        doc,
        [
            "LangGraph frente a un solo prompt: separa razonamiento, uso de tools y generación estructurada en nodos con aristas condicionales.",
            "ChromaDB local: vector store persistente sin servicio externo; menú en menu.json re-indexable vía POST /api/rag/ingest.",
            "Simulación web como canal principal: demo estable sin depender de túneles (ngrok/Cloudflare) ni créditos Twilio.",
            "Memoria por transcript en BD: cada turno reconstruye mensajes HumanMessage/AIMessage para contexto multi-turno.",
            "Confirmación híbrida: el LLM marca is_complete y el backend detecta frases («confirmo», «sí») o botón manual.",
            "WebSocket broadcast: desacopla la UI del polling; todos los paneles ven cambios al instante.",
            "Pydantic para contratos: órdenes tipadas entre agente, API y frontend.",
        ],
    )

    # 7
    add_heading(doc, "7. Limitaciones actuales", 1)
    add_bullets(
        doc,
        [
            "Dependencia de API OpenAI (LLM, embeddings, STT, TTS); sin clave no hay demo de IA.",
            "Twilio + túnel local (ngrok/Cloudflare): en pruebas desde Guatemala, Twilio no siempre alcanza el webhook del PC; la simulación web es la vía fiable.",
            "Precios y disponibilidad dependen de menu.json; no hay inventario en tiempo real externo.",
            "Un solo restaurante y menú fijo; sin multi-sucursal ni autenticación de usuarios.",
            "El agente puede equivocarse en ítems o confirmación; se requiere supervisión en demo.",
            "SQLite y Chroma en disco local; no preparado para alta concurrencia en producción.",
            "TTS/STT con latencia de red; no optimizado para llamadas masivas.",
        ],
    )

    # 8
    add_heading(doc, "8. Posibles mejoras futuras", 1)
    add_bullets(
        doc,
        [
            "Desplegar backend en la nube (Render, Railway, Azure) para webhooks Twilio estables.",
            "Autenticación y roles (agente, cocina, caja, administrador).",
            "Dashboard de métricas: tiempo de atención, pedidos por hora, errores del agente.",
            "Integración con POS o impresora de cocina (tickets automáticos).",
            "LangGraph checkpoints y persistencia de estado del grafo entre reinicios.",
            "Pruebas automatizadas (pytest, Playwright) y CI/CD.",
            "Soporte multi-idioma y menús dinámicos por temporada.",
            "Media Streams de Twilio para audio en tiempo real sobre WebSocket.",
        ],
    )

    # 9
    add_heading(doc, "9. High Level System Design", 1)
    add_para(doc, "Vista de componentes y flujo de información del sistema:", bold=True)

    hlsd = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                         CAPA DE PRESENTACIÓN (React)                         │
│  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────────┐  │
│  │ Agente       │ │ Cocina       │ │ Caja         │ │ Todos los pedidos│  │
│  │ (simulación) │ │              │ │              │ │                  │  │
│  └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └────────┬─────────┘  │
│         │                │                │                   │            │
│         └────────────────┴────────────────┴───────────────────┘            │
│                                    │ WebSocket + REST                       │
└────────────────────────────────────┼────────────────────────────────────────┘
                                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                      CAPA DE APLICACIÓN (FastAPI)                            │
│  ┌─────────────┐  ┌──────────────┐  ┌─────────────┐  ┌──────────────────┐  │
│  │ API REST    │  │ WebSocket    │  │ Twilio      │  │ Lifespan / CORS  │  │
│  │ /api/*      │  │ /ws/orders   │  │ /twilio/*   │  │                  │  │
│  └──────┬──────┘  └──────┬───────┘  └──────┬──────┘  └──────────────────┘  │
└─────────┼────────────────┼─────────────────┼──────────────────────────────┘
          │                │                 │
          ▼                ▼                 ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                         CAPA DE SERVICIOS                                    │
│  ┌──────────────────┐  ┌─────────────────┐  ┌────────────────────────────┐ │
│  │ LangGraph Agent  │  │ Order Service   │  │ Voice (STT/TTS OpenAI)     │ │
│  │ agent→tools→     │  │ sesiones,       │  │                            │ │
│  │ structure        │  │ estados pedido  │  │                            │ │
│  └────────┬─────────┘  └────────┬────────┘  └────────────────────────────┘ │
│           │                     │                                            │
│  ┌────────┴─────────┐  ┌───────┴────────┐  ┌────────────────────────────┐ │
│  │ Tools (LangChain)  │  │ RAG / ChromaDB │  │ OpenAI API                 │ │
│  │ buscar_en_menu,    │  │ embeddings     │  │ GPT-4o-mini, Whisper, TTS  │ │
│  │ validar, total     │  │ menu.json      │  │                            │ │
│  └────────────────────┘  └────────────────┘  └────────────────────────────┘ │
└─────────────────────────────────────────────────────────────────────────────┘
                                     │
                                     ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                         CAPA DE PERSISTENCIA                                 │
│              SQLite (call_sessions, orders, order_items)                     │
│              ChromaDB (vectores del menú en chroma_data/)                    │
└─────────────────────────────────────────────────────────────────────────────┘

Flujo resumido de una llamada simulada:
  1. Cliente habla/escribe → 2. STT (si audio) → 3. LangGraph + RAG/tools
  → 4. StructuredOrder → 5. Guardar en SQLite → 6. WebSocket → 7. UI Cocina/Caja
"""
    p = doc.add_paragraph()
    run = p.add_run(hlsd.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8)

    add_heading(doc, "9.1 Estructura del repositorio", 2)
    repo = """
Aplicacion web de IA/
├── backend/
│   ├── app/
│   │   ├── api/          # routes.py, twilio_routes.py
│   │   ├── models/       # SQLAlchemy
│   │   ├── schemas/      # Pydantic
│   │   ├── services/
│   │   │   ├── agent/    # LangGraph + tools
│   │   │   ├── rag/      # ChromaDB
│   │   │   └── voice/    # STT/TTS
│   │   └── websocket/
│   ├── data/menu.json
│   └── requirements.txt
├── frontend/
│   └── src/components/   # Agente, Cocina, Caja, pedidos
├── Documentacion.docx
└── README.md
"""
    p2 = doc.add_paragraph()
    r2 = p2.add_run(repo.strip())
    r2.font.name = "Consolas"
    r2.font.size = Pt(9)

    add_heading(doc, "10. Instalación y ejecución (resumen)", 1)
    add_bullets(
        doc,
        [
            "Backend: cd backend → venv → pip install -r requirements.txt → configurar .env → uvicorn app.main:app --reload --port 8000",
            "Frontend: cd frontend → npm install → npm run dev → http://localhost:5173",
            "Variable obligatoria: OPENAI_API_KEY en backend/.env",
            "Twilio opcional: credenciales + PUBLIC_BASE_URL pública + webhook de voz",
        ],
    )

    add_para(
        doc,
        "Integrantes: Diego Ernesto Merida Esteban (2890-22-18999), "
        "Cristian Alejandro Virula Hidalgo (2890-22-13840), "
        "Maycolm Anthony Gómez Pérez (2890-22-5945).",
    )


def main() -> None:
    doc = Document(str(DOC_PATH))
    trim_after_cover(doc)
    append_documentation(doc)
    doc.save(str(DOC_PATH))
    print(f"Documentación actualizada: {DOC_PATH}")


if __name__ == "__main__":
    main()
